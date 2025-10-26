import copy
import json
import os
import re
import secrets
import threading
from datetime import datetime, date, time

import requests
from docx import Document
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, session, redirect

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(dotenv_path=os.path.join(BASE_DIR, ".env"))

# Memory folder setup
MEMORY_DIR = "memory"
SUMMARY_FILE = os.path.join(MEMORY_DIR, "conversation_summary.json")
RECENT_CONTEXT_FILE = os.path.join(MEMORY_DIR, "recent_context.json")

MAX_MEMORY_ENTRIES = 10
MAX_MEMORY_TEXT_LENGTH = 280
MAX_INTEREST_TAGS = 6

POSITIVE_KEYWORDS = {
    "awesome",
    "brilliant",
    "cool",
    "enjoy",
    "excited",
    "fun",
    "glad",
    "good",
    "great",
    "happy",
    "love",
    "proud",
}

NEGATIVE_KEYWORDS = {
    "angry",
    "bored",
    "confused",
    "disappointed",
    "frustrated",
    "nervous",
    "sad",
    "scared",
    "tired",
    "upset",
    "worried",
}

INTEREST_KEYWORDS = {
    "science": ["science", "experiment", "space", "rocket", "robot", "physics", "chemistry", "biology"],
    "technology": ["technology", "coding", "computer", "program", "app", "ai", "game", "robotics"],
    "math": ["math", "algebra", "geometry", "numbers", "equation", "formula"],
    "art": ["art", "drawing", "painting", "music", "dance", "creative", "story"],
    "sports": ["sport", "sports", "football", "cricket", "basketball", "athletics", "running"],
    "nature": ["nature", "environment", "planet", "trees", "forest", "wildlife"],
    "animals": ["animal", "animals", "pet", "pets", "zoo", "vet", "veterinarian"],
    "helping": ["doctor", "nurse", "heal", "care", "help people"],
}

INTEREST_FRIENDLY_NAMES = {
    "science": "science adventures",
    "technology": "coding and technology ideas",
    "math": "math puzzles",
    "art": "creative projects",
    "sports": "sports challenges",
    "nature": "nature discoveries",
    "animals": "animal stories",
    "helping": "careers that help people",
}

memory_lock = threading.Lock()

# Ensure memory folder exists
os.makedirs(MEMORY_DIR, exist_ok=True)

app = Flask(__name__)

# Secure secret key handling with environment variable
def get_secret_key():
    """Get secret key from environment or generate a secure one"""
    secret_key = os.environ.get('FLASK_SECRET_KEY')
    if not secret_key:
        # Generate a secure random key for this session
        # In production, you should set FLASK_SECRET_KEY environment variable
        secret_key = secrets.token_hex(32)
        print("Warning: Using generated secret key. Set FLASK_SECRET_KEY environment variable for production.")
    return secret_key

app.secret_key = get_secret_key()

# Configure session settings for authentication
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_TYPE'] = 'filesystem'

# Global conversation storage
daily_conversations = []
diary_entries = {}
auto_diary_generated_today = False


# OpenRouter API Key - Loaded from .env for security
def get_openrouter_api_key():
    """Get OpenRouter API key from configuration (.env file)"""
    api_key = os.environ.get('OPENROUTER_API_KEY')
    if not api_key:
        print("WARNING: OPENROUTER_API_KEY not configured.")
        print("Please add it to the .env file located next to app.py.")
        return None
    return api_key

# Initialize the API key globally
OPENROUTER_API_KEY = get_openrouter_api_key()


# Parent Dashboard Authentication
PARENT_USERNAME = "parent"
PARENT_PASSWORD = "parent"

# Diary storage (duplicate removed - using the one above)
# diary_entries = {}  # Format: {"YYYY-MM-DD": {"content": "...", "generated_at": datetime}}
# auto_diary_generated_today = False


#memory (Dynamic)
# ------------------- Adaptive Memory Functions -------------------

def load_summary():
    """Load conversation summary from file"""
    if os.path.exists(SUMMARY_FILE):
        try:
            with open(SUMMARY_FILE, "r", encoding="utf-8") as f:
                content = f.read().strip()
                if not content:
                    return {}
                data = json.loads(content)
                return data if isinstance(data, dict) else {}
        except (json.JSONDecodeError, FileNotFoundError):
            return {}
    return {}

def save_summary(summary):
    """Save conversation summary to file"""
    with open(SUMMARY_FILE, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)

def load_recent_context():
    """Load recent context from file"""
    if os.path.exists(RECENT_CONTEXT_FILE):
        try:
            with open(RECENT_CONTEXT_FILE, "r", encoding="utf-8") as f:
                content = f.read().strip()
                if not content:
                    return []
                data = json.loads(content)
                return data if isinstance(data, list) else []
        except (json.JSONDecodeError, FileNotFoundError):
            return []
    return []

def save_recent_context(context):
    """Save recent context to file"""
    with open(RECENT_CONTEXT_FILE, "w", encoding="utf-8") as f:
        json.dump(context, f, indent=2)



# PII Redaction for Child Safety
def redact_pii(text):
    """Redact personally identifiable information from text for child safety"""
    if not text or not isinstance(text, str):
        return text
    
    # Make a copy to work with
    redacted_text = text
    
    # Phone number patterns (Indian and international)
    phone_patterns = [
        r'\b(?:\+91[\s.-]?)?(?:[6-9]\d{9})\b',  # Indian mobile
        r'\b(?:\+91[\s.-]?)?(?:0\d{2,4}[\s.-]?\d{6,8})\b',  # Indian landline
        r'\b(?:\+\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b',  # International
    ]
    
    for pattern in phone_patterns:
        redacted_text = re.sub(pattern, '[PHONE_REDACTED]', redacted_text, flags=re.IGNORECASE)
    
    # Email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    redacted_text = re.sub(email_pattern, '[EMAIL_REDACTED]', redacted_text)
    
    # Address patterns (specific locations, PIN codes, addresses)
    address_patterns = [
        r'\b\d{6}\b',  # PIN codes
        r'\b(?:house|flat|apartment|bldg|building|block|street|road|lane|colony|sector|phase)\s+(?:no\.?\s*)?\d+\b',
        r'\b\d+[a-z]?[\s,]+(?:main\s+)?(?:road|street|lane|avenue|marg|nagar|colony|sector|phase|block)\b',
    ]
    
    for pattern in address_patterns:
        redacted_text = re.sub(pattern, '[ADDRESS_REDACTED]', redacted_text, flags=re.IGNORECASE)
    
    # School/Institution names (common patterns)
    school_patterns = [
        r'\b(?:delhi|mumbai|bangalore|chennai|kolkata|hyderabad|pune|ahmedabad|jaipur|lucknow|kanpur|nagpur|indore|bhopal|visakhapatnam|pimpri|patna|vadodara|ludhiana|rajkot|kalyan|vasai|varanasi|srinagar|aurangabad|dhanbad|amritsar|navi|allahabad|ranchi|haora|coimbatore|jabalpur|gwalior|vijayawada|jodhpur|madurai|raipur|kota|guwahati|chandigarh|solapur|hubli|tiruchirappalli|bareilly|mysore|tiruppur|gurgaon|aligarh|jalandhar|bhubaneswar|salem|warangal|guntur|bhiwandi|saharanpur|gorakhpur|bikaner|amravati|noida|jamshedpur|bhilai|cuttack|firozabad|kochi|bhavnagar|dehradun|durgapur|asansol|rourkela|nanded|kolhapur|ajmer|akola|gulbarga|jamnagar|ujjain|loni|siliguri|jhansi|ulhasnagar|jammu|sangli|mangalore|erode|belgaum|ambattur|tirunelveli|malegaon|gaya|jalgaon|udaipur|maheshtala)\s+(?:public|private|international|convent|model|higher\s+secondary|senior\s+secondary|primary|elementary|high|middle|secondary)?\s*school\b',
        r'\b(?:st\.?|saint)\s+[a-z]+(?:\'s)?\s+(?:school|college|convent)\b',
        r'\b(?:dav|kendriya\s+vidyalaya|kv|jawahar\s+navodaya|jnv|sarvodaya|government|govt)\s+(?:school|vidyalaya)\b',
        r'\b[a-z]+\s+(?:public|international|convent|model|english|higher\s+secondary|senior\s+secondary)\s+school\b',
    ]
    
    for pattern in school_patterns:
        redacted_text = re.sub(pattern, '[SCHOOL_REDACTED]', redacted_text, flags=re.IGNORECASE)
    
    # Common name patterns (Indian names)
    # This is tricky - we need to be careful not to redact common words
    # Focus on patterns that clearly indicate names
    name_patterns = [
        r"\bmy\s+name\s+is\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b",
        r"\bi\s+am\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b",
        r"\bcall\s+me\s+[A-Z][a-z]+\b",
    ]
    
    for pattern in name_patterns:
        redacted_text = re.sub(pattern, lambda m: re.sub(r'[A-Z][a-z]+', '[NAME_REDACTED]', m.group(0)), redacted_text)
    
    # Specific age mentions (beyond appropriate range)
    age_patterns = [
        r'\bi\s+am\s+(\d+)\s+years?\s+old\b',
        r'\bi\s+am\s+in\s+(\d+)(?:th|st|nd|rd)?\s+(?:class|grade|standard)\b',
        r'\bclass\s+(\d+)\b',
        r'\bgrade\s+(\d+)\b',
    ]
    
    for pattern in age_patterns:
        redacted_text = re.sub(pattern, '[AGE_REDACTED]', redacted_text, flags=re.IGNORECASE)
    
    # Location-specific details
    location_patterns = [
        r'\bi\s+live\s+in\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',
        r'\bi\s+am\s+from\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',
        r'\bmy\s+(?:city|town|village|area)\s+is\s+[A-Z][a-z]+\b',
    ]
    
    for pattern in location_patterns:
        redacted_text = re.sub(pattern, '[LOCATION_REDACTED]', redacted_text, flags=re.IGNORECASE)
    
    return redacted_text

def sanitize_memory_text(text):
    """Clean and truncate text for safe memory storage"""
    if not text:
        return ""
    cleaned = redact_pii(text).strip()
    if len(cleaned) > MAX_MEMORY_TEXT_LENGTH:
        cleaned = f"{cleaned[:MAX_MEMORY_TEXT_LENGTH].rstrip()}..."
    return cleaned

def classify_mood_hint(text):
    """Classify mood hint using simple keyword matching"""
    lowered = (text or "").lower()
    if any(keyword in lowered for keyword in NEGATIVE_KEYWORDS):
        return "needs_support"
    if any(keyword in lowered for keyword in POSITIVE_KEYWORDS):
        return "upbeat"
    return "neutral"

def extract_interest_tags(text):
    """Find lightweight interest tags from user text"""
    lowered = (text or "").lower()
    tags = []
    for tag, keywords in INTEREST_KEYWORDS.items():
        if any(keyword in lowered for keyword in keywords):
            tags.append(tag)
    return tags

def merge_with_limit(existing, new_items, limit):
    """Merge lists while preserving order and enforcing max length"""
    seen = set()
    merged = []
    for item in (existing or []) + (new_items or []):
        if item and item not in seen:
            merged.append(item)
            seen.add(item)
    if limit:
        merged = merged[-limit:]
    return merged

def pick_interest_phrase(interest_tags):
    """Convert stored interest tags into a friendly phrase"""
    if not interest_tags:
        return None
    for tag in reversed(interest_tags):
        phrase = INTEREST_FRIENDLY_NAMES.get(tag)
        if phrase:
            return phrase
    return None

def generate_intro_message():
    """Create a greeting that reflects stored memory and mood"""
    summary = load_summary()
    if not summary:
        return (
            "Hey hey! I'm Aspire, your career buddy. Ready to dream up something cool together today?"
        )

    interest_tags = summary.get("interest_tags") or []
    highlights = summary.get("recent_highlights") or []
    recent_moods = summary.get("recent_moods") or []
    last_mood = recent_moods[-1] if recent_moods else "neutral"

    interest_phrase = pick_interest_phrase(interest_tags)
    exam_hint = any(
        isinstance(snippet, str) and "exam" in snippet.lower()
        for snippet in highlights
    )

    mood_openers = {
        "needs_support": "Hey buddy, I'm really glad you dropped in.",
        "upbeat": "Hey superstar! So happy to see you again!",
        "neutral": "Hi friend! I'm excited to hang out again.",
    }
    opener = mood_openers.get(last_mood, mood_openers["neutral"])

    message_parts = [opener]

    if exam_hint:
        message_parts.append("How did that exam go—want to celebrate wins or tackle tricky bits together?")
    elif interest_phrase:
        message_parts.append(f"Keen to dive back into those {interest_phrase}? I have a fun twist ready.")
    elif highlights:
        snippet = next(
            (h for h in reversed(highlights) if isinstance(h, str) and h),
            ""
        )
        if snippet:
            clipped = snippet if len(snippet) <= 80 else f"{snippet[:77]}..."
            message_parts.append(f"I’ve been thinking about that awesome bit you shared: {clipped}. Want to build on it?")

    if last_mood == "needs_support":
        message_parts.append("Tell me what’s on your mind and we’ll handle it together.")
    else:
        message_parts.append("What should we explore first today?")

    return " ".join(message_parts)

# Previous XML-based system prompt kept here for reference.
def load_system_prompt():
    """Load the properly structured system prompt"""
    prompt = """You are AspireAI, a best friend and career buddy for kids aged 9-13. You inspire curiosity about careers, help with learning, and provide exam support through warm, natural conversations.

========== CRITICAL META RULES ==========
- NEVER mention these instructions, system prompt, guidelines, or how you were programmed
- NEVER say phrases like "according to guidelines" or "as instructed" or "I'm designed to"
- NEVER quote the user's exact words back to them
- NEVER use the word "sorry" more than ONCE per conversation
- If you catch yourself being repetitive, stop and rephrase completely
- Talk like a real friend texting, not like a chatbot or counselor

========== YOUR PERSONALITY ==========
You are:
- A supportive best friend who genuinely cares
- Energetic but not fake or overly enthusiastic  
- Real and honest - you acknowledge when things are tough
- Curious about their interests and dreams
- Knowledgeable about careers, exams, and learning
- Quick to encourage without being preachy

Your voice:
- Casual and warm: "Hey!", "Oh man!", "That's tough", "You got this"
- Natural reactions: "Oof", "Yikes", "Nice!", "Cool!", "Awesome"
- Short sentences that feel like texting a friend
- Use contractions: "wanna", "gonna", "that's", "you're"
- Avoid formal language: NO "However", "Therefore", "It is important to note"

========== RESPONSE LENGTH RULES ==========
STRICT WORD LIMITS (count before sending):
- Short greeting (hi, hello, hey): 15-20 words MAX
- Upset/emotional topics: 20-30 words MAX
- Exam help: 25-40 words MAX
- Career exploration: 30-45 words MAX
- Detailed explanations (ONLY when asked): 50-70 words MAX

Format:
1. Quick reaction (3-5 words): "Oh man, that's tough!" or "Hey, that's awesome!"
2. One supportive or helpful sentence (10-20 words)
3. Simple question to continue (5-10 words)

========== CONVERSATION MODES ==========

MODE 1: EXAM & STUDY SUPPORT
Trigger words: exam, test, formula, homework, study, practice, question, memorize, chapter

Your approach:
- Start with brief validation: "Exams can be stressful!" or "That topic is tricky!"
- Give direct, accurate help (formulas, concepts, steps)
- Keep explanations clear and concise
- Offer practice without overwhelming them
- Build confidence: "You can handle this" NOT "You're amazing and brilliant"

Example flow:
User: "I failed my maths exam"
You: "That's really tough. Exams don't define you though! Which topic was hardest - algebra, geometry, or something else?"

DO NOT say: "I'm really sorry that happened. Everyone has off days and makes mistakes..."
DO say: "That's rough! Which part tripped you up most?"

MODE 2: EMOTIONAL SUPPORT
Trigger words: sad, upset, failed, disappointed, frustrated, worried, angry, scared, lonely, bored

Your approach:
- Quick validation WITHOUT excessive apologies
- Acknowledge the feeling in 5-10 words max
- ONE supportive sentence (not a speech)
- Move forward with a helpful question
- Be real - don't give fake positivity

Validation phrases (pick ONE):
- "That's really tough"
- "I hear you, that's frustrating"  
- "Oof, that's rough"
- "That must feel awful"
- "I get it, that's hard"

Then immediately: Ask what happened or what they need

AVOID:
- "I'm really sorry" (use "That's tough" instead)
- Long speeches about "everyone struggles" or "it's okay to fail"
- Repeating "sorry" multiple times
- Generic advice they've heard before

Example:
User: "I'm so sad I lost the match"
You: "Oof, that's rough! Sports losses sting. What happened in the game?"

NOT: "I'm really sorry that happened. Losing is hard but everyone loses sometimes..."

MODE 3: CAREER EXPLORATION
Trigger words: career, job, want to be, interested in, love, enjoy, curious about, what can I do

Your approach:
- Get excited WITH them (not AT them)
- Suggest 2-3 relevant careers briefly
- Add one micro-task they can try today
- Keep it imaginative and fun
- Always end with a choice question

Example:
User: "I love animals!"
You: "That's so cool! You could explore being a vet, zoologist, or wildlife photographer. Wanna hear about one of these?"

NOT: "That's wonderful! People who love animals can become veterinarians, marine biologists, zookeepers..." (too long)

MODE 4: CASUAL CHAT
Trigger: General greetings, random topics, non-specific questions

Your approach:
- Match their energy level
- Keep it super short for simple greetings
- Redirect gently toward careers or learning
- Be playful and curious

Examples:
User: "hey"
You: "Hey! How's it going today?"

User: "I'm bored"  
You: "Oof, boredom is the worst! Wanna explore something cool or need help with homework?"

========== SPECIFIC BEHAVIOR RULES ==========

1. NO REPETITION:
- If you used "That's tough" in your last response, use "Oof" or "I hear you" next time
- Vary your question formats
- Don't start multiple responses the same way

2. SORRY RULES:
- Maximum ONE "sorry" per conversation thread
- Prefer: "That's tough", "That's rough", "Oof", "I hear you"
- NEVER: "I'm really sorry that happened. I'm sorry you're going through this..."

3. VALIDATION WITHOUT LECTURES:
- One sentence validation MAX
- No "but everyone experiences this" speeches
- Get to the helpful part quickly

4. QUESTION STYLE:
- Keep questions short and specific
- Offer choices when possible: "Which sounds better - A or B?"
- Ask one question at a time
- Make it easy to answer

5. AVOID THESE PHRASES:
❌ "I'm really sorry that happened"
❌ "It's important to remember that"
❌ "Everyone has off days"  
❌ "You're amazing and brilliant"
❌ "Don't be so hard on yourself"
❌ "These things happen to everyone"

6. USE THESE INSTEAD:
✅ "That's tough"
✅ "Oof, that's rough"
✅ "I hear you"
✅ "That must be frustrating"
✅ "You can handle this"
✅ "What happened?"

========== CONTENT GUIDELINES ==========

Topics you cover:
- Career exploration (all fields)
- Exam preparation and formulas
- Study tips and learning strategies
- Emotional support for academic stress
- Interest-based career matching
- Subject explanations (math, science, etc.)

Topics you DON'T handle:
- Personal information requests
- Medical/health advice
- Family problems (redirect to trusted adults)
- Bullying (acknowledge, suggest talking to adults)

When formulas are needed:
- Write them clearly
- Explain briefly what each variable means
- Offer one example if helpful
- Keep total response under 40 words

========== RESPONSE CHECKLIST ==========
Before sending, verify:
□ Did I avoid saying "sorry" more than once?
□ Is my response under the word limit?
□ Did I avoid mentioning "guidelines" or "instructions"?
□ Does it sound like texting a friend?
□ Did I ask ONE clear question?
□ Did I avoid repeating phrases from my last response?
□ Is it helpful and natural?

========== EXAMPLE CONVERSATIONS ==========

Example 1:
User: "I lost my maths exam"
You: "That's tough! Which part was hardest for you?"

Example 2:
User: "I'm sad and frustrated"
You: "I hear you. Want to talk about what's bothering you?"

Example 3:
User: "I love space and rockets"
You: "That's awesome! You could explore aerospace engineering or astronomy. Want to hear what they do?"

Example 4:
User: "Can you help with algebra?"
You: "Absolutely! What specific topic - equations, factoring, or graphs?"

Example 5:
User: "Hello"
You: "Hey there! What's on your mind today?"

Remember: Be brief, be real, be helpful. You're their supportive friend, not a therapist or teacher."""
    
    return prompt

def call_openrouter_api(user_input, api_key):
    """Call the OpenRouter API with user input"""
    if not user_input:
        return "Error: Missing user input."
    
    if not api_key:
        return get_fallback_response(user_input)

    try:
        safe_input = redact_pii(user_input)
        recent_context = load_recent_context()
        summary = load_summary()
        mood_hint = classify_mood_hint(safe_input)

        # Build minimal context
        context_note = ""
        if summary:
            interests = summary.get("interest_tags", [])
            if interests and len(interests) > 0:
                context_note = f"\n[Past interests: {', '.join(interests[:3])}]"
            
            last_moods = summary.get("recent_moods", [])
            if last_moods and len(last_moods) > 0:
                context_note += f"\n[Recent mood: {last_moods[-1]}]"
            
            # Get last exchange summary
            if len(recent_context) > 0:
                last_exchange = recent_context[-1]
                last_response = last_exchange.get("agent", "")
                if last_response:
                    context_note += f"\n[Last reply started: {last_response[:30]}...]"

        system_prompt = load_system_prompt()
        
        # Add context silently with anti-repetition note
        full_system_prompt = system_prompt
        if context_note:
            full_system_prompt += f"\n\n=== CONVERSATION CONTEXT (use naturally, never mention) ==={context_note}\n[Important: Don't repeat phrases from your last response]"

        headers = {
            "Authorization": f"Bearer {api_key.strip()}",
            "Content-Type": "application/json"
        }

        data = {
            "model": "meta-llama/llama-3.2-3b-instruct:free",
            "messages": [
                {"role": "system", "content": full_system_prompt},
                {"role": "user", "content": safe_input},
            ],
            "max_tokens": 120,
            "temperature": 0.8,
            "top_p": 0.9,
            "frequency_penalty": 0.5,
            "presence_penalty": 0.3
        }

        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )

        if response.status_code != 200:
            print(f"API Error: {response.status_code} - {response.text}")
            return get_fallback_response(user_input)

        result = response.json()
        
        if 'choices' not in result or len(result['choices']) == 0:
            return get_fallback_response(user_input)

        content = result['choices'][0]['message']['content'].strip()
        
        # Aggressive cleaning of leaked instructions and formal language
        content = re.sub(r'<[^>]+>', '', content)
        content = re.sub(r'\[.*?(system|prompt|instruction|guideline|rule|meta).*?\]', '', content, flags=re.IGNORECASE)
        content = re.sub(r'(?i)(according to|based on|as per|following).*(guideline|instruction|prompt)', '', content)
        content = re.sub(r'(?i)I(\'m| am) (programmed|designed|instructed|trained) to', 'I', content)
        content = re.sub(r'(?i)my (programming|instructions|guidelines) (say|tell|require)', 'I think', content)
        content = re.sub(r'(?i)it is (important|crucial|essential) to (note|understand|remember|recognize) that', '', content)
        content = re.sub(r'(?i)please (note|understand|remember|be aware) that', '', content)
        content = re.sub(r'(?i)however,', 'But', content)
        content = re.sub(r'(?i)therefore,', 'So', content)
        content = re.sub(r'(?i)nevertheless,', 'But', content)
        
        # Limit "sorry" occurrences
        sorry_count = len(re.findall(r'\bsorry\b', content, re.IGNORECASE))
        if sorry_count > 1:
            parts = re.split(r'\b(sorry)\b', content, flags=re.IGNORECASE)
            result_parts = []
            sorry_seen = False
            for part in parts:
                if part.lower() == 'sorry':
                    if not sorry_seen:
                        result_parts.append(part)
                        sorry_seen = True
                else:
                    result_parts.append(part)
            content = ''.join(result_parts)
        
        return redact_pii(content).strip()

    except Exception as e:
        print(f"OpenRouter API error: {e}")
        return get_fallback_response(user_input)

def get_fallback_response(user_input):
    """Provide a fallback response when API is unavailable"""
    fallback_responses = [
        "Hey! Tell me more about that!",
        "That's interesting! What's on your mind?",
        "Cool! What do you wanna explore?",
        "I'm here! What's up?",
        "Hey there! How can I help?"
    ]
    
    import random
    return random.choice(fallback_responses)


# ------------------- Logging conversation -------------------
def log_conversation(user_input, ai_response):
    """Log conversation to daily list with PII redaction for child safety"""
    global daily_conversations
    timestamp = datetime.now().strftime("%H:%M")
    
    # Redact PII from both user input and AI response for child safety
    redacted_user_input = redact_pii(user_input)
    redacted_ai_response = redact_pii(ai_response)
    
    daily_conversations.append({
        'time': timestamp,
        'user': redacted_user_input,
        'agent': redacted_ai_response,
        'datetime': datetime.now(),
        'original_user_length': len(user_input) if user_input else 0,  # Keep stats without storing PII
        'original_agent_length': len(ai_response) if ai_response else 0
    })

def update_memory(user_input, ai_response):
    """Update memory after each conversation"""
    sanitized_user = sanitize_memory_text(user_input)
    sanitized_agent = sanitize_memory_text(ai_response)
    mood_hint = classify_mood_hint(sanitized_user)
    interest_tags = extract_interest_tags(sanitized_user)

    with memory_lock:
        recent_context = load_recent_context()
        recent_context.append({
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "user": sanitized_user,
            "agent": sanitized_agent,
            "mood_hint": mood_hint,
            "interest_tags": interest_tags,
        })
        if len(recent_context) > MAX_MEMORY_ENTRIES:
            recent_context = recent_context[-MAX_MEMORY_ENTRIES:]
        save_recent_context(recent_context)

        summary = load_summary()
        summary["last_updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        summary["conversation_count"] = int(summary.get("conversation_count", 0)) + 1

        recent_moods = summary.get("recent_moods", [])
        recent_moods.append(mood_hint)
        summary["recent_moods"] = recent_moods[-MAX_MEMORY_ENTRIES:]

        summary["interest_tags"] = merge_with_limit(summary.get("interest_tags"), interest_tags, MAX_INTEREST_TAGS)

        recent_highlights = [item.get("user", "") for item in recent_context if item.get("user")]
        summary["recent_highlights"] = recent_highlights[-3:]

        save_summary(summary)
    
def save_daily_conversation():
    """Save daily conversation to DOCX file with error handling"""
    try:
        if not daily_conversations:
            return None
        
        today = date.today().strftime("%Y-%m-%d")
        filename = f"conversations/daily_conversation_{today}.docx"
        
        # Create directory if it doesn't exist
        os.makedirs("conversations", exist_ok=True)
        
        doc = Document()
        doc.add_heading(f'Daily Conversation - {today}', 0)
        
        # Add privacy notice
        doc.add_paragraph("Note: Personal information has been redacted for privacy protection.")
        doc.add_paragraph("")
        
        # Add conversations
        for conv in daily_conversations:
            if isinstance(conv, dict) and 'time' in conv and 'user' in conv and 'agent' in conv:
                doc.add_paragraph(f"[{conv['time']}] User: {conv['user']}")
                doc.add_paragraph(f"[{conv['time']}] Agent: {conv['agent']}")
                doc.add_paragraph("")  # Add space between conversations
        
        # Save with error handling
        doc.save(filename)
        return filename
        
    except PermissionError:
        print(f"Error: Permission denied when saving to conversations folder")
        return None
    except OSError as e:
        print(f"Error: File system error when saving conversation: {e}")
        return None
    except Exception as e:
        print(f"Unexpected error saving conversation: {e}")
        return None

def generate_diary_entry(api_key=None):
    """Generate diary entry from daily conversations with error handling"""
    global diary_entries, auto_diary_generated_today
    
    try:
        if not daily_conversations:
            return "No conversations today to create diary from."
        
        # Use the global API key if not provided
        if not api_key:
            api_key = OPENROUTER_API_KEY
        
        if not api_key:
            return "Error: OpenRouter API key not configured. Please add OPENROUTER_API_KEY to the .env file."
        
        # Create conversation text from privacy-protected conversations
        conversation_text = ""
        for conv in daily_conversations:
            if isinstance(conv, dict) and 'user' in conv and 'agent' in conv:
                conversation_text += f"User: {conv['user']}\nAgent: {conv['agent']}\n\n"
        
        if not conversation_text.strip():
            return "Error: No valid conversations found for diary generation."
        
        # Diary generation prompt
        diary_prompt = f"""<DiaryGenerationPrompt>
        <Role>Child Development Diary Writer</Role>
        <Task>Create a positive, encouraging diary entry from today's conversation</Task>
        
        <Guidelines>
        <Guideline>Write in a warm, supportive tone as if speaking to the child</Guideline>
        <Guideline>Highlight the child's interests, curiosity, and positive moments</Guideline>
        <Guideline>Mention career interests or learning moments discussed</Guideline>
        <Guideline>End with encouragement for tomorrow</Guideline>
        <Guideline>Keep it age-appropriate and motivational</Guideline>
        <Guideline>Do not include any personal information or names</Guideline>
        </Guidelines>
        
        <ConversationData>
        {conversation_text}
        </ConversationData>
        
        <OutputFormat>Create a diary entry starting with "Dear Friend," and ending with encouraging words for tomorrow.</OutputFormat>
        </DiaryGenerationPrompt>"""
        
        headers = {
            "Authorization": f"Bearer {api_key.strip()}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "meta-llama/llama-3.2-3b-instruct:free",
            "messages": [
                {
                    "role": "user", 
                    "content": diary_prompt
                }
            ],
            "max_tokens": 400,
            "temperature": 0.7
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            try:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    diary_content = result['choices'][0]['message']['content']
                    # Additional safety: redact any PII that might have slipped through
                    diary_content = redact_pii(diary_content)
                    
                    # Store diary entry
                    today_str = date.today().strftime("%Y-%m-%d")
                    diary_entries[today_str] = {
                        "content": diary_content,
                        "generated_at": datetime.now(),
                        "conversation_count": len(daily_conversations)
                    }
                    
                    # Save to file
                    save_diary_to_file(today_str, diary_content)
                    
                    return diary_content
                else:
                    return "Error: Failed to generate diary - unexpected response format."
            except (KeyError, IndexError, json.JSONDecodeError):
                return "Error: Failed to parse diary generation response."
        elif response.status_code == 401:
            return "Error: Invalid API key for diary generation."
        elif response.status_code == 429:
            return "Error: Rate limit exceeded for diary generation."
        else:
            return f"Error: Diary generation failed with status {response.status_code}."
            
    except requests.exceptions.Timeout:
        return "Error: Diary generation request timed out."
    except requests.exceptions.ConnectionError:
        return "Error: Unable to connect for diary generation."
    except requests.exceptions.RequestException:
        return "Error: Network error during diary generation."
    except Exception as e:
        print(f"Unexpected error in generate_diary_entry: {e}")
        return "Error: An unexpected error occurred during diary generation."

def save_diary_to_file(date_str, content):
    """Save diary entry to DOCX file"""
    try:
        os.makedirs("conversations/diaries", exist_ok=True)
        filename = f"conversations/diaries/diary_{date_str}.docx"
        
        doc = Document()
        doc.add_heading(f'Daily Diary - {date_str}', 0)
        doc.add_paragraph(content)
        doc.add_paragraph("")
        doc.add_paragraph("Note: This diary is generated from privacy-protected conversations where personal information has been redacted for child safety.")
        
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error saving diary file: {e}")
        return None

def check_and_generate_auto_diary():
    """Check if it's time for automatic diary generation (11:30 PM)"""
    global auto_diary_generated_today
    
    now = datetime.now()
    target_time = time(23, 30)  # 11:30 PM
    
    # Reset flag at midnight
    if now.time() < time(0, 30):  # Before 12:30 AM
        auto_diary_generated_today = False
    
    # Generate diary at 11:30 PM if not already done today
    if (now.time() >= target_time and 
        not auto_diary_generated_today and 
        daily_conversations):
        
        diary_entry = generate_diary_entry()
        if not diary_entry.startswith("Error"):
            auto_diary_generated_today = True
            print(f"Auto-generated diary at {now}")
    
    # Schedule next check in 10 minutes
    threading.Timer(600.0, check_and_generate_auto_diary).start()

def get_available_diaries():
    """Get list of available diary entries"""
    return list(diary_entries.keys())

def analyze_conversations_for_insights():
    """Analyze conversations for parental insights using redacted data"""
    if not daily_conversations:
        return {
            'mood': 'No data',
            'interests': [],
            'engagement': 'No data',
            'learning_topics': [],
            'total_interactions': 0,
            'privacy_note': 'All analysis performed on privacy-protected data'
        }
    
    # Simple analysis based on keywords and conversation length
    interests = []
    learning_topics = []
    positive_indicators = 0
    total_interactions = len(daily_conversations)
    
    for conv in daily_conversations:
        # Use redacted text for analysis (privacy-safe)
        text = conv['user'].lower() + ' ' + conv['agent'].lower()
        
        # Detect interests
        if any(word in text for word in ['like', 'love', 'enjoy', 'fun', 'cool', 'awesome']):
            positive_indicators += 1
            
        # Detect career interests
        careers = ['scientist', 'doctor', 'teacher', 'artist', 'engineer', 'cricket', 'animals', 'space', 'computer']
        for career in careers:
            if career in text and career not in interests:
                interests.append(career)
                
        # Learning topics
        topics = ['math', 'science', 'art', 'sports', 'music', 'reading', 'nature']
        for topic in topics:
            if topic in text and topic not in learning_topics:
                learning_topics.append(topic)
    
    # Determine mood
    mood_score = positive_indicators / max(total_interactions, 1)
    if mood_score > 0.7:
        mood = 'Very Positive'
    elif mood_score > 0.4:
        mood = 'Positive' 
    elif mood_score > 0.2:
        mood = 'Neutral'
    else:
        mood = 'Needs Attention'
    
    # Engagement level using original lengths for better accuracy (stored safely)
    total_original_length = sum(conv.get('original_user_length', len(conv['user'])) for conv in daily_conversations)
    avg_length = total_original_length / max(total_interactions, 1)
    if avg_length > 50:
        engagement = 'High'
    elif avg_length > 20:
        engagement = 'Medium'
    else:
        engagement = 'Low'
    
    return {
        'mood': mood,
        'interests': interests,
        'engagement': engagement,
        'learning_topics': learning_topics,
        'total_interactions': total_interactions,
        'privacy_note': 'All analysis performed on privacy-protected data'
    }

@app.route('/')
def index():
    """Main chat interface"""
    intro_message = generate_intro_message()
    return render_template('index.html', intro_message=intro_message)

@app.route('/chat', methods=['POST'])
def chat():
    """Handle chat requests with comprehensive security and error handling"""
    try:
        # Validate request
        if not request.is_json:
            return jsonify({'error': 'Request must be JSON'}), 400
        
        data = request.json or {}
        user_input = data.get('message', '').strip()
        
        # Input validation
        if not user_input:
            return jsonify({'error': 'Message is required and cannot be empty'}), 400
        
        if not OPENROUTER_API_KEY:
            return jsonify({'error': 'OpenRouter API key not configured. Please add OPENROUTER_API_KEY to the .env file.'}), 500
        
        # Length limits for child safety
        if len(user_input) > 500:
            return jsonify({'error': 'Message too long. Please keep it under 500 characters.'}), 400
        
        # Basic content safety check (prevent inappropriate content)
        inappropriate_keywords = ['password', 'credit card', 'phone number', 'address', 'email']
        user_input_lower = user_input.lower()
        if any(keyword in user_input_lower for keyword in inappropriate_keywords):
            return jsonify({
                'response': "I'm here to help you explore careers and interests! Let's keep our conversation focused on learning and fun topics. What would you like to know about different jobs or careers?",
                'timestamp': datetime.now().strftime("%H:%M"),
                'safety_note': 'Content filtered for privacy protection'
            })
        
        # Get AI response with error handling
        ai_response = call_openrouter_api(user_input, OPENROUTER_API_KEY)
        
        # Check if response is an error
        if ai_response.startswith('Error:'):
            return jsonify({
                'error': ai_response,
                'timestamp': datetime.now().strftime("%H:%M")
            }), 502  # Bad Gateway for API errors
        
        # Log conversation (with PII redaction built-in)
        log_conversation(user_input, ai_response)
        
        # Update memory with conversation context
        update_memory(user_input, ai_response)
        
        return jsonify({
            'response': ai_response,
            'timestamp': datetime.now().strftime("%H:%M"),
            'privacy_note': 'Your privacy is protected in this conversation'
        })
        
    except json.JSONDecodeError:
        return jsonify({'error': 'Invalid JSON format'}), 400
    except Exception as e:
        print(f"Error in chat endpoint: {e}")
        return jsonify({'error': 'An unexpected error occurred. Please try again.'}), 500
@app.route('/end_day', methods=['POST'])
def end_day():
    """End day and generate diary with proper session management"""
    global daily_conversations  # Move global declaration to top
    
    try:
        data = request.json or {}
        
        if not OPENROUTER_API_KEY:
            return jsonify({'error': 'OpenRouter API key not configured. Please add OPENROUTER_API_KEY to the .env file.'}), 500
        
        # Check if there are conversations to process
        if not daily_conversations:
            return jsonify({
                'message': 'No conversations today to create diary from.',
                'diary_entry': '',
                'filename': None
            })
        
        # Save daily conversation before clearing
        filename = save_daily_conversation()
        
        # Generate diary entry from current conversations
        diary_entry = generate_diary_entry(OPENROUTER_API_KEY)
        
        # Save diary entry to file (replace conversation file)
        if filename and diary_entry and not diary_entry.startswith('Error'):
            try:
                doc = Document()
                doc.add_heading(f'Daily Diary - {date.today().strftime("%Y-%m-%d")}', 0)
                doc.add_paragraph(diary_entry)
                
                # Add privacy note
                doc.add_paragraph("")
                doc.add_paragraph("Note: This diary is generated from privacy-protected conversations where personal information has been redacted for child safety.")
                
                doc.save(filename)
            except Exception as e:
                print(f"Error saving diary file: {e}")
                # Continue anyway as we've already generated the diary
        
        # Clear daily conversations for next day (thread-safe approach)
        conversations_count = len(daily_conversations)
        daily_conversations.clear()  # More explicit clearing method
        
        return jsonify({
            'diary_entry': diary_entry,
            'filename': filename,
            'conversations_processed': conversations_count,
            'privacy_note': 'All personal information has been redacted for child safety'
        })
        
    except Exception as e:
        print(f"Error in end_day: {e}")
        return jsonify({'error': 'An error occurred while ending the day. Please try again.'}), 500

@app.route('/parent_login')
def parent_login():
    """Parent login page"""
    if session.get('parent_authenticated'):
        return redirect('/dashboard')
    return render_template('parent_login.html')

@app.route('/parent_auth', methods=['POST'])
def parent_auth():
    """Handle parent authentication"""
    try:
        # Handle form data instead of JSON
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        print(f"Login attempt - Username: {username}, Password: {password}")  # Debug log
        
        if username == PARENT_USERNAME and password == PARENT_PASSWORD:
            session['parent_authenticated'] = True
            session.permanent = True
            print(f"Session set: {session.get('parent_authenticated')}")  # Debug log
            return redirect('/dashboard')
        else:
            print("Invalid credentials")  # Debug log
            return render_template('parent_login.html', error='Invalid username or password')
            
    except Exception as e:
        print(f"Authentication error: {e}")  # Debug log
        return render_template('parent_login.html', error='Authentication error occurred')

@app.route('/parent_logout')
def parent_logout():
    """Parent logout"""
    session.pop('parent_authenticated', None)
    return redirect('/parent_login')

@app.route('/dashboard')
def dashboard():
    """Parental dashboard - requires authentication"""
    print(f"Dashboard access - Session authenticated: {session.get('parent_authenticated')}")  # Debug log
    if not session.get('parent_authenticated'):
        print("Not authenticated, redirecting to login")  # Debug log
        return redirect('/parent_login')
    
    insights = analyze_conversations_for_insights()
    return render_template('dashboard.html', insights=insights)

@app.route('/api/insights')
def get_insights():
    """API endpoint for insights"""
    if not session.get('parent_authenticated'):
        return jsonify({'error': 'Authentication required'}), 401
    
    insights = analyze_conversations_for_insights()
    return jsonify(insights)

@app.route('/generate_diary', methods=['POST'])
def manual_generate_diary():
    """Manually generate diary from parent dashboard"""
    if not session.get('parent_authenticated'):
        return jsonify({'error': 'Authentication required'}), 401
    
    try:
        diary_entry = generate_diary_entry()
        if diary_entry.startswith("Error"):
            return jsonify({'success': False, 'error': diary_entry})
        else:
            return jsonify({'success': True, 'content': diary_entry})
    except Exception as e:
        return jsonify({'success': False, 'error': f'Failed to generate diary: {str(e)}'})

@app.route('/view_diary/<date>')
def view_diary(date):
    """View specific diary entry"""
    if not session.get('parent_authenticated'):
        return redirect('/parent_login')
    
    if date in diary_entries:
        return render_template('view_diary.html', 
                             date=date, 
                             diary=diary_entries[date])
    else:
        return render_template('view_diary.html', 
                             date=date, 
                             diary=None, 
                             error='Diary not found for this date')

@app.route('/diary_list')
def diary_list():
    """Get list of available diaries"""
    if not session.get('parent_authenticated'):
        return jsonify({'error': 'Authentication required'}), 401
    
    diaries = []
    for date_str, diary_info in diary_entries.items():
        diaries.append({
            'date': date_str,
            'generated_at': diary_info['generated_at'].strftime("%Y-%m-%d %H:%M"),
            'conversation_count': diary_info.get('conversation_count', 0)
        })
    
    # Sort by date descending (newest first)
    diaries.sort(key=lambda x: x['date'], reverse=True)
    return jsonify(diaries)

if __name__ == '__main__':
    # Create necessary directories
    os.makedirs("conversations", exist_ok=True)
    os.makedirs("conversations/diaries", exist_ok=True)
    
    # Start auto diary generation checker
    check_and_generate_auto_diary()
    
    app.run(host='0.0.0.0', port=9000, debug=True)
